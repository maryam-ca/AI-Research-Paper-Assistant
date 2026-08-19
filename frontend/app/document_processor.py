"""Document processing: extract text from PDF/DOCX and chunk it."""
import io
import re
import logging
from pypdf import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger("document_processor")


def extract_text_from_pdf(file_bytes: bytes, return_page_indices: bool = True):
    """Return list of {"text": str, "page": int} (1-indexed)."""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001
            logger.warning("Failed to extract page %s: %s", i, e)
            text = ""
        if text.strip():
            pages.append({"text": text, "page": i})
    return pages


def extract_text_from_docx(file_bytes: bytes):
    """DOCX has no reliable page numbers; we return page=1 for all blocks."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return [{"text": text, "page": 1}] if text.strip() else []


def chunk_text(full_text: str, chunk_size: int = 1000, overlap: int = 200):
    """Split text into overlapping chunks by character count.

    Keeps chunks coherent by respecting sentence boundaries where possible.
    """
    if not full_text or not full_text.strip():
        return []

    sentences = re.split(r"(?<=[.!?])\s+", full_text.strip())
    chunks = []
    current = ""
    for sent in sentences:
        if len(current) + len(sent) + 1 <= chunk_size:
            current = (current + " " + sent).strip() if current else sent
        else:
            if current:
                chunks.append(current)
            # start new chunk, carrying overlap from the tail of the previous
            if len(sent) > chunk_size:
                # Very long sentence: hard split
                for i in range(0, len(sent), chunk_size - overlap):
                    chunks.append(sent[i : i + chunk_size])
                current = ""
            else:
                current = sent
    if current:
        chunks.append(current)

    # Apply overlap by prepending tail of previous chunk
    overlapped = []
    for idx, c in enumerate(chunks):
        if idx > 0 and overlap > 0:
            prev = chunks[idx - 1]
            c = (prev[-overlap:] + " " + c).strip()
        overlapped.append(c)
    return overlapped


def process_document(file_bytes: bytes, source_type: str):
    """Unified entry: returns (pages, full_text)."""
    if source_type == "pdf":
        pages = extract_text_from_pdf(file_bytes)
    elif source_type in ("docx", "url", "doi"):
        # url/doi fetched content is treated as docx-style text blob
        pages = extract_text_from_docx(file_bytes)
    else:
        pages = extract_text_from_pdf(file_bytes)
    full_text = "\n".join(p["text"] for p in pages)
    return pages, full_text

